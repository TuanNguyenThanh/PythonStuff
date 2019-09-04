from docx import Document
from docx.shared import Inches
from simplify_docx import simplify
# from docx import Section

document = Document('test.docx')
paragraphs = document.paragraphs
tables = document.tables
sections = document.sections

for section in sections:
    print(section.footer.paragraphs[0].text)
    print(section.header.paragraphs[0].text)

for table in tables:
    # print(len(table.rows))
    # print(len(table.columns))
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

# print(len(paragraphs))
# for p in paragraphs:
#     if p.text != '': print(p.text)
#     else: print('No')
# print(paragraphs[13].text)

# def findObject(objectname):
